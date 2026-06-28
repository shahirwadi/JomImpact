from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = "JomImpact_Use_Case_Catalog.docx"
DIAGRAM_OUTPUT = "JomImpact_Class_Diagram.png"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
TEXT = RGBColor(30, 30, 30)
MUTED = RGBColor(85, 85, 85)
HEADER_FILL = "E8EEF5"
BORDER = "CBD5E1"


def load_font(size, bold=False):
    candidates = [
        "arialbd.ttf" if bold else "arial.ttf",
        "calibrib.ttf" if bold else "calibri.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def wrap_text(draw, text, font, max_width):
    words = text.split()
    lines = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_arrow(draw, start, end, label="", color=(71, 85, 105)):
    draw.line([start, end], fill=color, width=3)
    x1, y1 = start
    x2, y2 = end
    dx = x2 - x1
    dy = y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 14
    points = [
        (x2, y2),
        (x2 - ux * size + px * size * 0.45, y2 - uy * size + py * size * 0.45),
        (x2 - ux * size - px * size * 0.45, y2 - uy * size - py * size * 0.45),
    ]
    draw.polygon(points, fill=color)
    if label:
        font = load_font(22)
        mx = (x1 + x2) / 2
        my = (y1 + y2) / 2
        bbox = draw.textbbox((0, 0), label, font=font)
        pad = 5
        draw.rounded_rectangle(
            [mx - (bbox[2] / 2) - pad, my - 14, mx + (bbox[2] / 2) + pad, my + 16],
            radius=6,
            fill=(255, 255, 255),
            outline=(226, 232, 240),
        )
        draw.text((mx - bbox[2] / 2, my - 12), label, fill=color, font=font)


def draw_label(draw, x, y, label, color=(71, 85, 105)):
    font = load_font(22)
    bbox = draw.textbbox((0, 0), label, font=font)
    pad = 5
    draw.rounded_rectangle(
        [x - pad, y - pad, x + bbox[2] + pad, y + bbox[3] + pad],
        radius=6,
        fill=(255, 255, 255),
        outline=(226, 232, 240),
    )
    draw.text((x, y), label, fill=color, font=font)


def draw_polyline_arrow(draw, points, label="", label_at=None, color=(71, 85, 105)):
    draw.line(points, fill=color, width=3, joint="curve")
    x1, y1 = points[-2]
    x2, y2 = points[-1]
    dx = x2 - x1
    dy = y2 - y1
    length = max((dx * dx + dy * dy) ** 0.5, 1)
    ux, uy = dx / length, dy / length
    px, py = -uy, ux
    size = 14
    arrow = [
        (x2, y2),
        (x2 - ux * size + px * size * 0.45, y2 - uy * size + py * size * 0.45),
        (x2 - ux * size - px * size * 0.45, y2 - uy * size - py * size * 0.45),
    ]
    draw.polygon(arrow, fill=color)
    if label and label_at:
        draw_label(draw, label_at[0], label_at[1], label, color=color)


def draw_class_box(draw, x, y, w, h, name, fields, methods=None, fill=(248, 250, 252)):
    methods = methods or []
    border = (37, 99, 155)
    header = (232, 238, 245)
    text = (30, 41, 59)
    muted = (71, 85, 105)
    title_font = load_font(26, bold=True)
    body_font = load_font(22)
    method_font = load_font(21)
    draw.rounded_rectangle([x, y, x + w, y + h], radius=12, fill=fill, outline=border, width=3)
    draw.rounded_rectangle([x, y, x + w, y + 48], radius=12, fill=header, outline=border, width=3)
    draw.line([(x, y + 48), (x + w, y + 48)], fill=border, width=2)
    title_bbox = draw.textbbox((0, 0), name, font=title_font)
    draw.text((x + (w - title_bbox[2]) / 2, y + 10), name, fill=(15, 60, 102), font=title_font)

    cursor = y + 62
    for field in fields:
        for line in wrap_text(draw, field, body_font, w - 30):
            draw.text((x + 16, cursor), line, fill=text, font=body_font)
            cursor += 27
    if methods:
        draw.line([(x + 12, cursor + 4), (x + w - 12, cursor + 4)], fill=(203, 213, 225), width=2)
        cursor += 14
        for method in methods:
            for line in wrap_text(draw, method, method_font, w - 30):
                draw.text((x + 16, cursor), line, fill=muted, font=method_font)
                cursor += 25


def create_class_diagram(path=DIAGRAM_OUTPUT):
    img = Image.new("RGB", (2200, 1450), "white")
    draw = ImageDraw.Draw(img)
    title_font = load_font(36, bold=True)
    subtitle_font = load_font(24)
    draw.text((70, 40), "JomImpact Core Domain Class Diagram", fill=(11, 37, 69), font=title_font)
    draw.text(
        (70, 88),
        "Main model classes and Firestore document relationships represented in the current system.",
        fill=(85, 85, 85),
        font=subtitle_font,
    )

    boxes = {
        "UserModel": (70, 160, 500, 330),
        "EventModel": (850, 160, 560, 310),
        "ApplicationModel": (790, 610, 650, 370),
        "ImpactAward": (1570, 650, 520, 270),
        "FeedPostModel": (70, 700, 520, 300),
        "FeedCommentModel": (70, 1110, 520, 230),
        "MarketplaceItemModel": (850, 1110, 560, 270),
        "MarketplacePurchaseModel": (1570, 1110, 520, 250),
    }

    draw_class_box(
        draw,
        *boxes["UserModel"],
        "UserModel",
        [
            "+ id, name, email, photoUrl",
            "+ role: UserRole",
            "+ location, state, skills",
            "+ organization, totalEvents, totalHours",
            "+ organizerApprovalStatus",
        ],
        ["+ isOrganizerApproved", "+ copyWith()", "+ toMap()/fromMap()"],
    )
    draw_class_box(
        draw,
        *boxes["EventModel"],
        "EventModel",
        [
            "+ id, organizerId, organizerName",
            "+ title, description, location, state",
            "+ startDate, endDate, category, status",
            "+ maxVolunteers, currentVolunteers",
            "+ requirements, benefits, imageUrl",
        ],
        ["+ isFull, spotsLeft, fillRate", "+ copyWith()", "+ toMap()/fromMap()"],
    )
    draw_class_box(
        draw,
        *boxes["ApplicationModel"],
        "ApplicationModel",
        [
            "+ id, eventId, eventTitle",
            "+ volunteerId, volunteerName, volunteerBio",
            "+ status, message, reviewNotes",
            "+ attendanceStatus, verifiedHours",
            "+ impactPoints, appliedAt",
        ],
        ["+ copyWith()", "+ toMap()/fromMap()"],
    )
    draw_class_box(
        draw,
        *boxes["ImpactAward"],
        "ImpactAward",
        [
            "<<Firestore document>>",
            "+ eventId, eventTitle",
            "+ applicationId, volunteerId",
            "+ hours, points",
            "+ awardedAt, awardedBy",
        ],
        ["points = hours x 10"],
        fill=(255, 251, 235),
    )
    draw_class_box(
        draw,
        *boxes["FeedPostModel"],
        "FeedPostModel",
        [
            "+ id, authorId, authorName",
            "+ authorRole, content, imageUrl",
            "+ likedBy, commentCount",
            "+ createdAt, updatedAt",
        ],
        ["+ likeCount", "+ isLikedBy()", "+ canEdit"],
    )
    draw_class_box(
        draw,
        *boxes["FeedCommentModel"],
        "FeedCommentModel",
        [
            "+ id, postId",
            "+ authorId, authorName, authorRole",
            "+ content, createdAt",
        ],
        ["+ toMap()/fromMap()"],
    )
    draw_class_box(
        draw,
        *boxes["MarketplaceItemModel"],
        "MarketplaceItemModel",
        [
            "+ id, organizerId, organizerName",
            "+ title, description, price, imageUrl",
            "+ status, adminNotes",
            "+ reviewedBy, reviewedAt, createdAt",
        ],
        ["+ toMap()/fromMap()"],
    )
    draw_class_box(
        draw,
        *boxes["MarketplacePurchaseModel"],
        "MarketplacePurchaseModel",
        [
            "+ id, itemId, itemTitle",
            "+ organizerId, buyerId, buyerName",
            "+ price, status, createdAt",
        ],
        ["+ toMap()"],
    )

    draw_arrow(draw, (570, 260), (850, 260), "organizer owns")
    draw_arrow(draw, (1120, 470), (1120, 610), "has applications")
    draw_polyline_arrow(
        draw,
        [(570, 370), (660, 370), (660, 750), (790, 750)],
        "volunteer submits",
        (610, 535),
    )
    draw_arrow(draw, (1440, 800), (1570, 800), "finalizes into")
    draw_arrow(draw, (330, 490), (330, 700), "author")
    draw_arrow(draw, (330, 1000), (330, 1110), "has comments")
    draw_polyline_arrow(
        draw,
        [(570, 445), (700, 445), (700, 1210), (850, 1210)],
        "organizer lists",
        (635, 780),
    )
    draw_arrow(draw, (1410, 1235), (1570, 1235), "purchased as")

    legend_font = load_font(22)
    draw.rounded_rectangle([1500, 160, 2090, 410], radius=12, fill=(248, 250, 252), outline=(203, 213, 225), width=2)
    draw.text((1530, 188), "Legend", fill=(15, 60, 102), font=load_font(26, bold=True))
    legend_lines = [
        "Solid arrows show stored IDs or",
        "collection relationships.",
        "ImpactAward is represented because",
        "Firestore writes it during finalization,",
        "although no Dart model class exists.",
    ]
    y = 235
    for line in legend_lines:
        draw.text((1530, y), line, fill=(71, 85, 105), font=legend_font)
        y += 30

    img.save(path)
    return path


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in {
        "top": top,
        "start": start,
        "bottom": bottom,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_width(table, widths):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(width * 1440)))


def add_para(doc, text="", style=None, bold=False, italic=False, color=TEXT):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color
    return paragraph


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_borders(table)
    set_table_width(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = DARK_BLUE
    for row_data in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row_data):
            cell = cells[i]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.add_run(value)
    doc.add_paragraph()
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "right_margin", "bottom_margin", "left_margin"):
        setattr(section, side, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "JomImpact System Use Case Catalogue"
    header.runs[0].font.size = Pt(9)
    header.runs[0].font.color.rgb = MUTED
    footer = section.footer.paragraphs[0]
    footer.text = "Generated from Flutter, Firebase service, model, viewmodel, and Firestore rules inspection"
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].font.color.rgb = MUTED


def build_document():
    diagram_path = create_class_diagram()
    doc = Document()
    configure_styles(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("JomImpact System Use Case Catalogue")
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run(
        "Actors, functions, data relationships, and backend controls identified from the current Flutter/Firebase codebase."
    )
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED

    doc.add_heading("1. System Overview", level=1)
    add_para(
        doc,
        "JomImpact is a volunteering platform with three main roles: volunteer, organizer, and admin. "
        "The app uses Flutter with Provider viewmodels, Firebase Authentication for sign-in, Cloud Firestore for core data, and Cloudinary for image upload URLs."
    )
    add_table(
        doc,
        ["Layer", "Main files", "Responsibility"],
        [
            (
                "Presentation",
                "lib/views/auth, lib/views/volunteer, lib/views/organizer, lib/views/admin, lib/views/shared",
                "Screens and navigation for role-specific workflows.",
            ),
            (
                "State / orchestration",
                "AuthViewModel, EventViewModel, FeedViewModel, MarketplaceViewModel",
                "Holds UI state, loading/errors, filters, and calls service methods.",
            ),
            (
                "Domain models",
                "UserModel, EventModel, ApplicationModel, ImpactSummary, FeedPostModel, MarketplaceItemModel",
                "Defines fields, enums, serialization, and computed helpers.",
            ),
            (
                "Backend services",
                "FirebaseAuthService, FirebaseEventService, FirebaseFeedService, FirebaseMarketplaceService, CloudinaryImageService",
                "Performs Firebase Auth, Firestore reads/writes/transactions, and image uploads.",
            ),
            (
                "Security rules",
                "firestore.rules",
                "Enforces role permissions, ownership, approval status, field constraints, and immutable records.",
            ),
        ],
        [1.1, 2.3, 3.1],
    )

    doc.add_heading("2. Actors", level=1)
    add_table(
        doc,
        ["Actor", "Primary purpose", "Core relationships"],
        [
            (
                "Volunteer",
                "Finds opportunities, applies to events, tracks applications and impact, posts in feed, and buys approved marketplace items.",
                "Owns user profile; creates applications; receives impact awards; creates feed posts/comments; creates marketplace purchases.",
            ),
            (
                "Organizer",
                "Registers an organization account, waits for admin approval, creates events, reviews applicants, verifies attendance, finalizes impact, and submits marketplace items.",
                "Owns events; reviews applications for own events; awards points after attendance; owns marketplace listings.",
            ),
            (
                "Admin",
                "Approves or rejects organizer accounts and marketplace submissions.",
                "Updates organizer approval state; reviews marketplace item status.",
            ),
        ],
        [1.2, 2.4, 2.9],
    )

    doc.add_heading("3. Use Case Catalogue", level=1)
    use_cases = [
        (
            "UC-01",
            "Register account",
            "Volunteer or organizer",
            "Creates Firebase Auth user and matching users document. Organizer accounts start as pending; volunteers are immediately usable.",
            "FirebaseAuthService.register; AuthViewModel.register",
            "users",
        ),
        (
            "UC-02",
            "Log in and route by role",
            "All users",
            "Authenticates credentials, fetches the user profile, then AppRouter sends admins, approved organizers, pending organizers, or volunteers to the correct area.",
            "FirebaseAuthService.login/fetchCurrentUser; AuthViewModel._init/login; AppRouter",
            "users",
        ),
        (
            "UC-03",
            "Complete or edit profile",
            "All users",
            "Updates name, bio, phone, photo URL, location, state, skills, or organization while preserving role and approval state for non-admin users.",
            "FirebaseAuthService.updateUser; AuthViewModel.updateProfile",
            "users",
        ),
        (
            "UC-04",
            "Review organizer registration",
            "Admin",
            "Lists pending organizer accounts and sets approval status to approved or rejected with optional notes.",
            "getPendingOrganizerRequests; reviewOrganizerRequest",
            "users",
        ),
        (
            "UC-05",
            "Browse and filter events",
            "Volunteer",
            "Loads published events and supports search, category, and Malaysian state filtering.",
            "getAllPublishedEvents; EventViewModel.loadAllEvents/setSearchQuery/setCategory/setStateFilter",
            "events, users",
        ),
        (
            "UC-06",
            "View organizer directory and public profile",
            "Volunteer",
            "Shows approved organizers and their public details/events so volunteers can discover organizations.",
            "getAllOrganizers; getUserById; EventViewModel.loadOrganizers",
            "users, events",
        ),
        (
            "UC-07",
            "Apply for event",
            "Volunteer",
            "Creates one pending application for a published event if the event exists, is open, and has capacity.",
            "applyForEvent; EventViewModel.applyForEvent",
            "applications -> events/users",
        ),
        (
            "UC-08",
            "Track and withdraw application",
            "Volunteer",
            "Loads the volunteer's applications and allows withdrawal while the event is still active; accepted withdrawals reduce event capacity.",
            "getApplicationsForVolunteer; withdrawApplication; volunteerApplicationsStream",
            "applications, events",
        ),
        (
            "UC-09",
            "Create event",
            "Approved organizer",
            "Publishes an event with category, date range, capacity, image URL, requirements, benefits, location, and state. Increments organizer totalEvents.",
            "createEvent; EventViewModel.createEvent",
            "events -> users",
        ),
        (
            "UC-10",
            "Manage own events",
            "Approved organizer",
            "Lists, updates, and deletes organizer-owned events; delete also removes related applications and decrements totalEvents.",
            "getOrganizerEvents; updateEvent; deleteEvent; organizerEventsStream",
            "events, applications, users",
        ),
        (
            "UC-11",
            "Review volunteer applications",
            "Approved organizer",
            "Moves applications through pending, reviewing, accepted, waitlisted, or rejected. Accepted status changes keep currentVolunteers in sync.",
            "getApplicationsForEvent; updateApplicationStatus",
            "applications, applicationReviews, events",
        ),
        (
            "UC-12",
            "Complete event",
            "Approved organizer",
            "Marks an event completed only after the end time and before finalization.",
            "markEventCompleted",
            "events",
        ),
        (
            "UC-13",
            "Review attendance",
            "Approved organizer",
            "For accepted applications on completed events, records attended, partial, noShow, or excused plus verified hours.",
            "reviewAttendance",
            "applications",
        ),
        (
            "UC-14",
            "Finalize impact awards",
            "Approved organizer",
            "Requires attendance review for every accepted volunteer, writes impact points as verifiedHours x 10, creates impactAwards, and finalizes the event.",
            "finalizeEvent; impactSummaryStream",
            "events, applications, impactAwards",
        ),
        (
            "UC-15",
            "View impact summary",
            "Volunteer",
            "Aggregates impact awards into points, hours, events, badge tier, and progress toward the next tier.",
            "ImpactSummary; impactSummaryStream",
            "impactAwards",
        ),
        (
            "UC-16",
            "Create social feed post",
            "Organizer or volunteer",
            "Publishes text and/or image URL to the public role feed with author snapshot fields.",
            "FirebaseFeedService.createPost; FeedViewModel.createPost",
            "feedPosts, users",
        ),
        (
            "UC-17",
            "Interact with feed",
            "Organizer or volunteer",
            "Streams posts and comments, adds comments, toggles likes, and tracks comment count.",
            "postsStream; commentsStream; addComment; toggleLike",
            "feedPosts/comments",
        ),
        (
            "UC-18",
            "Edit or delete own feed post",
            "Organizer or volunteer",
            "Service supports post update and delete; model exposes a 10-minute edit window helper.",
            "updatePost; deletePost; FeedPostModel.canEdit",
            "feedPosts",
        ),
        (
            "UC-19",
            "Submit marketplace item",
            "Approved organizer",
            "Creates a pending item with title, description, price, optional image URL, and organizer ownership.",
            "FirebaseMarketplaceService.createItem; MarketplaceViewModel.createItem",
            "marketplaceItems",
        ),
        (
            "UC-20",
            "Review marketplace item",
            "Admin",
            "Approves or rejects pending items and stores admin notes, reviewedBy, and reviewedAt.",
            "reviewItem; pendingItemsStream",
            "marketplaceItems",
        ),
        (
            "UC-21",
            "Browse and buy marketplace item",
            "Volunteer",
            "Streams approved marketplace items and creates a pendingPayment purchase record using item and buyer snapshots.",
            "approvedItemsStream; createPurchase; MarketplaceViewModel.buyItem",
            "marketplaceItems -> marketplacePurchases",
        ),
        (
            "UC-22",
            "Upload image",
            "Organizer or volunteer",
            "Uploads selected file to Cloudinary and stores only the returned secure URL in app records.",
            "CloudinaryImageService.uploadImage",
            "Cloudinary, users/events/feedPosts/marketplaceItems imageUrl fields",
        ),
    ]
    add_table(
        doc,
        ["ID", "Use case", "Actor", "Function explanation", "Key functions", "Data"],
        use_cases,
        [0.55, 1.05, 1.0, 2.0, 1.25, 0.65],
    )

    doc.add_heading("4. Class Diagram", level=1)
    add_para(
        doc,
        "The diagram below focuses on the core domain model classes and Firestore-backed document relationships used by the system. "
        "Viewmodels and services sit above these models and are summarized in the System Overview table."
    )
    diagram_paragraph = doc.add_paragraph()
    diagram_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    diagram_run = diagram_paragraph.add_run()
    diagram_run.add_picture(diagram_path, width=Inches(6.5))
    add_table(
        doc,
        ["Class / document", "Purpose", "Important relationships"],
        [
            (
                "UserModel",
                "Represents volunteer, organizer, or admin profile data plus approval/location state.",
                "Organizer users own EventModel and MarketplaceItemModel records; volunteer users submit ApplicationModel records, create purchases, and receive impact awards.",
            ),
            (
                "EventModel",
                "Represents a volunteer opportunity with organizer ownership, schedule, category, capacity, and lifecycle status.",
                "Has many ApplicationModel records. Completed events can be finalized into ImpactAward documents.",
            ),
            (
                "ApplicationModel",
                "Represents a volunteer's application to one event, including review status, attendance status, verified hours, and points.",
                "Joins UserModel and EventModel. Organizer review notes live in applicationReviews; finalized attendance can create impactAwards.",
            ),
            (
                "ImpactAward",
                "Firestore document written at finalization time to summarize awarded hours and points.",
                "Connects finalized events, accepted applications, volunteers, and organizer-awarded impact points.",
            ),
            (
                "FeedPostModel / FeedCommentModel",
                "Represent the social feed and nested comments with author snapshot data.",
                "FeedPostModel has many FeedCommentModel records; both reference the author's user ID and role.",
            ),
            (
                "MarketplaceItemModel / MarketplacePurchaseModel",
                "Represent organizer-submitted marketplace listings and volunteer purchase intents.",
                "Approved MarketplaceItemModel records can be copied into MarketplacePurchaseModel records by volunteer buyers.",
            ),
        ],
        [1.55, 2.35, 2.6],
    )

    doc.add_heading("5. Main Data Relationships", level=1)
    add_table(
        doc,
        ["Relationship", "Meaning", "Rules / behavior"],
        [
            (
                "users -> events",
                "An approved organizer owns each event through organizerId.",
                "Only approved organizers may create events; only the owning organizer may update or delete.",
            ),
            (
                "users -> applications -> events",
                "A volunteer applies to an event through an application document.",
                "Volunteer can read own applications; organizer can read applications for owned events.",
            ),
            (
                "applications -> applicationReviews",
                "Organizer review notes are stored separately from public application state.",
                "Only the owning event organizer can create/read the private review note document.",
            ),
            (
                "events -> applications -> impactAwards",
                "Completed events lead to attendance review; finalized accepted applications can create impact awards.",
                "Points must equal verified hours x 10; noShow/excused records must have zero hours.",
            ),
            (
                "users -> feedPosts -> comments",
                "Organizers and volunteers create public community posts and comments.",
                "Posts/comments store author snapshots; likes and comment counts are update-only fields.",
            ),
            (
                "users -> marketplaceItems -> marketplacePurchases",
                "Organizers submit items, admins approve them, volunteers create purchase records.",
                "Volunteers can buy only approved items; purchase values must match the approved item snapshot.",
            ),
        ],
        [1.5, 2.6, 2.4],
    )

    doc.add_heading("6. Permission and State Controls", level=1)
    add_bullets(
        doc,
        [
            "Authentication is required for all Firestore reads and writes in the active rules.",
            "Users can create only their own volunteer or organizer account. Admin creation is outside the UI and seeded manually.",
            "Non-admin users cannot promote themselves or change organizer approval status.",
            "Approved organizer status is required for event creation, application review, attendance review, finalization, and marketplace item submission.",
            "Application review is closed once an event is completed, finalized, or cancelled, except attendance review after completion.",
            "Deletes are generally blocked in Firestore rules for user-facing records; some service methods include delete calls that require matching deployed rules to allow them.",
            "Malaysian location and state validation appears in auth/profile, event creation/update, and Firestore rules.",
        ],
    )

    doc.add_heading("7. Navigation Map", level=1)
    add_table(
        doc,
        ["Entry condition", "Destination", "Purpose"],
        [
            ("Not logged in", "LoginScreen", "Start authentication or registration."),
            ("Admin", "AdminMain", "Review organizer requests, marketplace requests, and profile."),
            (
                "Organizer pending or rejected",
                "OrganizerApprovalScreen",
                "Show approval status before organizer features are available.",
            ),
            (
                "Missing location/state",
                "LocationSetupScreen",
                "Require location setup before main role area.",
            ),
            ("Approved organizer", "OrganizerMain", "Dashboard, events, marketplace, feed, profile."),
            ("Volunteer", "VolunteerMain", "Browse events, organizers, applications, marketplace, feed, profile."),
        ],
        [1.6, 1.7, 3.2],
    )

    doc.add_heading("8. Notes and Gaps Observed", level=1)
    add_bullets(
        doc,
        [
            "Firestore rules disallow feed post delete and most deletes, while FirebaseFeedService.deletePost and FirebaseEventService.deleteEvent exist. If those features are visible, the deployed rules may block them.",
            "FeedPostModel exposes canEdit, but Firestore rules shown only allow likedBy and commentCount updates. Post text/image edits may be blocked unless rules are expanded.",
            "Marketplace purchases are created with pendingPayment status; no payment settlement or fulfillment workflow is implemented in the inspected code.",
            "Admin accounts are not self-registered in the app and must be seeded or created manually.",
        ],
    )

    doc.save(OUTPUT)


if __name__ == "__main__":
    build_document()
